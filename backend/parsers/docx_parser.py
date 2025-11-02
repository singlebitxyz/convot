"""
DOCX Parser

Extracts text from Microsoft Word DOCX documents using python-docx.
Preserves paragraph structure and handles formatting.
"""

from typing import Optional
import io
import logging
from parsers.base import BaseParser, ParseResult

logger = logging.getLogger(__name__)


class DOCXParser(BaseParser):
    """Parser for DOCX documents"""
    
    def __init__(self):
        super().__init__()
        self._docx = None
    
    def _get_docx(self):
        """Lazy load python-docx to avoid import errors if not installed"""
        if self._docx is None:
            try:
                from docx import Document
                self._docx = Document
            except ImportError:
                raise ImportError(
                    "python-docx is required for DOCX parsing. "
                    "Install it with: pip install python-docx"
                )
        return self._docx
    
    def parse(self, file_content: bytes, file_path: Optional[str] = None) -> ParseResult:
        """
        Extract text from DOCX file.
        
        Args:
            file_content: DOCX file content as bytes
            file_path: Optional file path for error messages
        
        Returns:
            ParseResult with extracted text
        """
        try:
            Document = self._get_docx()
            
            # Create file-like object from bytes
            docx_file = io.BytesIO(file_content)
            
            # Load document
            doc = Document(docx_file)
            
            text_parts = []
            metadata = {
                "paragraph_count": 0,
                "total_chars": 0,
            }
            
            # Extract text from all paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
                    metadata["paragraph_count"] += 1
                    metadata["total_chars"] += len(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())
                            metadata["total_chars"] += len(cell.text)
            
            full_text = "\n\n".join(text_parts)
            
            if not full_text.strip():
                return ParseResult(
                    text="",
                    metadata=metadata,
                    success=False,
                    error_message="DOCX contains no extractable text."
                )
            
            metadata["extracted_chars"] = len(full_text)
            
            self.logger.info(
                f"Successfully parsed DOCX: {metadata['paragraph_count']} paragraphs, "
                f"{metadata['extracted_chars']} characters"
            )
            
            return ParseResult(
                text=full_text,
                metadata=metadata
            )
            
        except Exception as e:
            error_msg = f"Failed to parse DOCX: {str(e)}"
            self.logger.error(error_msg, exc_info=True)
            return ParseResult(
                text="",
                metadata={},
                success=False,
                error_message=error_msg
            )
    
    def can_parse(self, mime_type: str, file_extension: Optional[str] = None) -> bool:
        """Check if this parser can handle DOCX files"""
        return (
            mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or
            (file_extension and file_extension.lower() == ".docx")
        )
    
    def get_supported_types(self) -> list[str]:
        """Get supported MIME types"""
        return ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]

